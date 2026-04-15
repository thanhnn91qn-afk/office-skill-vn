#!/usr/bin/env python3
# -*- coding: utf-8 -*-
from __future__ import annotations

import argparse
import re
import shutil
import sys
from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING
from docx.enum.style import WD_STYLE_TYPE
from docx.oxml import OxmlElement, parse_xml
from docx.oxml.ns import qn
from docx.shared import Cm, Mm, Pt

TOTAL_CONTENT_WIDTH_TWIPS = 9354
LEFT_COL_TWIPS = 3400
RIGHT_COL_TWIPS = TOTAL_CONTENT_WIDTH_TWIPS - LEFT_COL_TWIPS
QUOC_HIEU = "CỘNG HÒA XÃ HỘI CHỦ NGHĨA VIỆT NAM"
VIETNAMESE_DIACRITICS = set("ăâđêôơưáàảãạấầẩẫậắằẳẵặéèẻẽẹếềểễệóòỏõọốồổỗộớờởỡợúùủũụứừửữựíìỉĩịýỳỷỹỵ")

# Embedded XML for Straight Connector paragraphs (agency cell and motto cell).
# Coordinates extracted from the real Mau_cong_van_ND30_tai_ve.docx template.
# These replace the external oxml_underline_after_motto.xml file.
_AGENCY_CONNECTOR_XML = (
    '<w:p xmlns:w="http://schemas.openxmlformats.org/wordprocessingml/2006/main"'
    ' xmlns:wp="http://schemas.openxmlformats.org/drawingml/2006/wordprocessingDrawing"'
    ' xmlns:a="http://schemas.openxmlformats.org/drawingml/2006/main"'
    ' xmlns:wps="http://schemas.microsoft.com/office/word/2010/wordprocessingShape">'
    '<w:pPr><w:spacing w:after="0" w:line="288" w:lineRule="auto"/>'
    '<w:jc w:val="center"/></w:pPr>'
    '<w:r><w:rPr><w:noProof/></w:rPr><w:drawing>'
    '<wp:anchor distT="0" distB="0" distL="114300" distR="114300" simplePos="0"'
    ' relativeHeight="251658752" behindDoc="0" locked="0" layoutInCell="1" allowOverlap="1">'
    '<wp:simplePos x="0" y="0"/>'
    '<wp:positionH relativeFrom="column"><wp:posOffset>553085</wp:posOffset></wp:positionH>'
    '<wp:positionV relativeFrom="paragraph"><wp:posOffset>23495</wp:posOffset></wp:positionV>'
    '<wp:extent cx="1781175" cy="0"/>'
    '<wp:effectExtent l="0" t="0" r="0" b="0"/>'
    '<wp:wrapNone/>'
    '<wp:docPr id="2" name="Straight Connector 2"/>'
    '<wp:cNvGraphicFramePr/>'
    '<a:graphic>'
    '<a:graphicData uri="http://schemas.microsoft.com/office/word/2010/wordprocessingShape">'
    '<wps:wsp><wps:cNvCnPr/>'
    '<wps:spPr>'
    '<a:xfrm><a:off x="0" y="0"/><a:ext cx="1781175" cy="0"/></a:xfrm>'
    '<a:prstGeom prst="line"><a:avLst/></a:prstGeom>'
    '</wps:spPr>'
    '<wps:style>'
    '<a:lnRef idx="1"><a:schemeClr val="dk1"/></a:lnRef>'
    '<a:fillRef idx="0"><a:schemeClr val="dk1"/></a:fillRef>'
    '<a:effectRef idx="0"><a:schemeClr val="dk1"/></a:effectRef>'
    '<a:fontRef idx="minor"><a:schemeClr val="tx1"/></a:fontRef>'
    '</wps:style>'
    '<wps:bodyPr/></wps:wsp>'
    '</a:graphicData></a:graphic>'
    '</wp:anchor></w:drawing></w:r></w:p>'
)

_MOTTO_CONNECTOR_XML = (
    '<w:p xmlns:w="http://schemas.openxmlformats.org/wordprocessingml/2006/main"'
    ' xmlns:wp="http://schemas.openxmlformats.org/drawingml/2006/wordprocessingDrawing"'
    ' xmlns:a="http://schemas.openxmlformats.org/drawingml/2006/main"'
    ' xmlns:wps="http://schemas.microsoft.com/office/word/2010/wordprocessingShape">'
    '<w:pPr><w:spacing w:after="0" w:line="288" w:lineRule="auto"/>'
    '<w:jc w:val="center"/></w:pPr>'
    '<w:r><w:rPr><w:noProof/></w:rPr><w:drawing>'
    '<wp:anchor distT="0" distB="0" distL="114300" distR="114300" simplePos="0"'
    ' relativeHeight="251656704" behindDoc="0" locked="0" layoutInCell="1" allowOverlap="1">'
    '<wp:simplePos x="0" y="0"/>'
    '<wp:positionH relativeFrom="column"><wp:posOffset>841375</wp:posOffset></wp:positionH>'
    '<wp:positionV relativeFrom="paragraph"><wp:posOffset>23495</wp:posOffset></wp:positionV>'
    '<wp:extent cx="1781175" cy="0"/>'
    '<wp:effectExtent l="0" t="0" r="0" b="0"/>'
    '<wp:wrapNone/>'
    '<wp:docPr id="1" name="Straight Connector 1"/>'
    '<wp:cNvGraphicFramePr/>'
    '<a:graphic>'
    '<a:graphicData uri="http://schemas.microsoft.com/office/word/2010/wordprocessingShape">'
    '<wps:wsp><wps:cNvCnPr/>'
    '<wps:spPr>'
    '<a:xfrm><a:off x="0" y="0"/><a:ext cx="1781175" cy="0"/></a:xfrm>'
    '<a:prstGeom prst="line"><a:avLst/></a:prstGeom>'
    '</wps:spPr>'
    '<wps:style>'
    '<a:lnRef idx="1"><a:schemeClr val="dk1"/></a:lnRef>'
    '<a:fillRef idx="0"><a:schemeClr val="dk1"/></a:fillRef>'
    '<a:effectRef idx="0"><a:schemeClr val="dk1"/></a:effectRef>'
    '<a:fontRef idx="minor"><a:schemeClr val="tx1"/></a:fontRef>'
    '</wps:style>'
    '<wps:bodyPr/></wps:wsp>'
    '</a:graphicData></a:graphic>'
    '</wp:anchor></w:drawing></w:r></w:p>'
)


def _set_run_font_tnr(run, *, size_pt: int | None = None, bold: bool | None = None, italic: bool | None = None) -> None:
    run.font.name = "Times New Roman"
    if size_pt is not None:
        run.font.size = Pt(size_pt)
    if bold is not None:
        run.font.bold = bold
    if italic is not None:
        run.font.italic = italic
    r = run._element
    rpr = r.get_or_add_rPr()
    rfonts = rpr.get_or_add_rFonts()
    rfonts.set(qn("w:ascii"), "Times New Roman")
    rfonts.set(qn("w:hAnsi"), "Times New Roman")
    rfonts.set(qn("w:eastAsia"), "Times New Roman")


def _set_para_with_br(p, text: str) -> None:
    """Set paragraph content, converting literal \n to <w:br/> (soft line return)."""
    p.clear()
    parts = text.split("\n")
    for i, part in enumerate(parts):
        if i > 0:
            br_run = p.add_run()
            br = OxmlElement("w:br")
            br_run._r.append(br)
        if part:
            p.add_run(part)


def _set_cell_width(cell, twips: int) -> None:
    tc = cell._tc
    tc_pr = tc.get_or_add_tcPr()
    for old in list(tc_pr):
        if old.tag == qn("w:tcW"):
            tc_pr.remove(old)
    tc_w = OxmlElement("w:tcW")
    tc_w.set(qn("w:type"), "dxa")
    tc_w.set(qn("w:w"), str(twips))
    tc_pr.append(tc_w)


def _table_no_borders(table) -> None:
    tbl = table._tbl
    tbl_pr = tbl.tblPr
    if tbl_pr is None:
        tbl_pr = OxmlElement("w:tblPr")
        tbl.insert(0, tbl_pr)
    tbl_borders = OxmlElement("w:tblBorders")
    for name in ("top", "left", "bottom", "right", "insideH", "insideV"):
        el = OxmlElement(f"w:{name}")
        el.set(qn("w:val"), "nil")
        el.set(qn("w:sz"), "0")
        el.set(qn("w:space"), "0")
        tbl_borders.append(el)
    tbl_pr.append(tbl_borders)


def _set_table_fixed_layout(table, total_width_twips: int) -> None:
    tbl = table._tbl
    tbl_pr = tbl.tblPr
    if tbl_pr is None:
        tbl_pr = OxmlElement("w:tblPr")
        tbl.insert(0, tbl_pr)

    for old in list(tbl_pr):
        if old.tag in (qn("w:tblW"), qn("w:tblLayout")):
            tbl_pr.remove(old)

    tbl_w = OxmlElement("w:tblW")
    tbl_w.set(qn("w:type"), "dxa")
    tbl_w.set(qn("w:w"), str(total_width_twips))
    tbl_pr.append(tbl_w)

    tbl_layout = OxmlElement("w:tblLayout")
    tbl_layout.set(qn("w:type"), "fixed")
    tbl_pr.append(tbl_layout)


def _set_cell_paragraph_lines(cell, lines: list[str]) -> None:
    if not lines:
        lines = [""]
    while len(cell.paragraphs) < len(lines):
        cell.add_paragraph("")
    for idx, line in enumerate(lines):
        cell.paragraphs[idx].text = line
    # Keep only the requested semantic lines (exclude shape paragraphs added later).
    while len(cell.paragraphs) > len(lines):
        last = cell.paragraphs[-1]
        if _paragraph_has_line_shape(last._element):
            break
        el = last._element
        el.getparent().remove(el)


def _build_underline_paragraph(kind: str):
    """
    Build a Straight Connector paragraph for the ND30 header cells.
    kind:
      - 'agency': line under issuing agency (left header cell)
      - 'motto': line under motto/tiêu ngữ (right header cell)
    XML coordinates are extracted from Mau_cong_van_ND30_tai_ve.docx directly.
    No external XML file is required.
    """
    if kind not in {"agency", "motto"}:
        raise ValueError("kind must be 'agency' or 'motto'")
    xml = _AGENCY_CONNECTOR_XML if kind == "agency" else _MOTTO_CONNECTOR_XML
    el = parse_xml(xml)
    # Keep line centered relative to the current header cell.
    for node in el.iter():
        if str(node.tag).endswith("positionH"):
            for c in list(node):
                node.remove(c)
            align = OxmlElement("wp:align")
            align.text = "center"
            node.append(align)
    return el


def _paragraph_has_line_shape(p_el) -> bool:
    for el in p_el.iter():
        if el.tag.endswith("}prstGeom") and el.get("prst") == "line":
            return True
    return False


def _cell_has_line_shape(cell) -> bool:
    tc = cell._tc
    for child in tc:
        if child.tag == qn("w:p") and _paragraph_has_line_shape(child):
            return True
    return False


def _assert_header_underlines(doc: Document) -> None:
    if len(doc.tables) < 1:
        raise SystemExit("Thiếu bảng header để kiểm tra gạch.")
    t0 = doc.tables[0]
    left0 = t0.rows[0].cells[0]
    right0 = t0.rows[0].cells[-1]
    if not _cell_has_line_shape(left0):
        raise SystemExit("Thiếu gạch dưới tên cơ quan.")
    if not _cell_has_line_shape(right0):
        raise SystemExit("Thiếu gạch dưới quốc hiệu/tiêu ngữ.")


def _remove_line_shape_paragraphs_from_cell(cell) -> None:
    tc = cell._tc
    for child in list(tc):
        if child.tag == qn("w:p") and _paragraph_has_line_shape(child):
            tc.remove(child)


def ensure_underline_after_motto_cell(cell) -> None:
    _remove_line_shape_paragraphs_from_cell(cell)
    new_el = _build_underline_paragraph("motto")
    tc = cell._tc
    w_ps = [c for c in tc if c.tag == qn("w:p")]
    insert_after = w_ps[1] if len(w_ps) >= 2 else (w_ps[0] if w_ps else None)
    if insert_after is not None:
        txt = "".join(insert_after.itertext())
        if "Độc lập" not in txt and "Hạnh phúc" not in txt:
            insert_after = None
    if insert_after is None:
        for p_el in reversed(w_ps):
            text = "".join(p_el.itertext())
            if "Độc lập" in text or "Hạnh phúc" in text:
                insert_after = p_el
                break
    if insert_after is None:
        for p_el in reversed(w_ps):
            if "".join(p_el.itertext()).strip():
                insert_after = p_el
                break
    if insert_after is None:
        tc.append(new_el)
        return
    children = list(tc)
    idx = children.index(insert_after)
    tc.insert(idx + 1, new_el)
    if not _cell_has_line_shape(cell):
        raise SystemExit("Khong the chen gach duoi cho quoc hieu - tieu ngu.")


def ensure_underline_after_agency_cell(cell) -> None:
    _remove_line_shape_paragraphs_from_cell(cell)
    new_el = _build_underline_paragraph("agency")
    tc = cell._tc
    w_ps = [c for c in tc if c.tag == qn("w:p")]
    # Place line right under agency-name line (prefer second paragraph).
    insert_after = w_ps[1] if len(w_ps) >= 2 else (w_ps[0] if w_ps else None)
    if insert_after is not None and not "".join(insert_after.itertext()).strip():
        insert_after = None
    if insert_after is None:
        for p_el in reversed(w_ps):
            if "".join(p_el.itertext()).strip():
                insert_after = p_el
                break
    if insert_after is None:
        tc.append(new_el)
        return
    children = list(tc)
    idx = children.index(insert_after)
    tc.insert(idx + 1, new_el)
    if not _cell_has_line_shape(cell):
        raise SystemExit("Khong the chen gach duoi cho ten co quan.")


def _remove_paragraph_elements_between(doc: Document, start_el, end_el) -> None:
    body = doc._body._body  # type: ignore[attr-defined]
    children = list(body.iterchildren())
    try:
        i_start = children.index(start_el)
        i_end = children.index(end_el)
    except ValueError:
        return
    if i_end <= i_start:
        return
    for el in children[i_start + 1 : i_end]:
        if el.tag == qn("w:p"):
            body.remove(el)


def _build_nd30_document_from_source(texts: list[str], *, justify_body: bool, body_pt: int) -> Document:
    doc = Document()
    _configure_normal_style(doc, body_pt)
    section = doc.sections[0]
    _set_section_a4_margins(section)

    # Keep header for page numbering only.
    header = section.header
    hp = header.paragraphs[0] if header.paragraphs else header.add_paragraph()
    hp.text = ""
    hp.alignment = WD_ALIGN_PARAGRAPH.CENTER
    hr = hp.add_run()
    _set_run_font_tnr(hr, size_pt=body_pt, bold=False, italic=False)
    _add_page_field(hr)

    # Header block is in document body (not Word header/footer).
    t0 = doc.add_table(rows=3, cols=2)
    _table_no_borders(t0)
    _set_table_fixed_layout(t0, TOTAL_CONTENT_WIDTH_TWIPS)
    t0.autofit = False
    _set_cell_width(t0.rows[0].cells[0], LEFT_COL_TWIPS)
    _set_cell_width(t0.rows[0].cells[1], RIGHT_COL_TWIPS)
    _set_cell_width(t0.rows[1].cells[0], LEFT_COL_TWIPS)
    _set_cell_width(t0.rows[1].cells[1], RIGHT_COL_TWIPS)

    left0 = t0.rows[0].cells[0]
    right0 = t0.rows[0].cells[1]
    left1 = t0.rows[1].cells[0]
    right1 = t0.rows[1].cells[1]

    _set_cell_paragraph_lines(left0, ["SỞ Y TẾ TỈNH QUẢNG NINH", "BỆNH VIỆN ĐA KHOA TỈNH QUẢNG NINH"])
    _ensure_agency_header_lines(left0)

    _set_cell_paragraph_lines(right0, [QUOC_HIEU, "Độc lập - Tự do - Hạnh phúc"])
    for p in right0.paragraphs:
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        for run in p.runs:
            _set_run_font_tnr(run, size_pt=13, bold=True, italic=False)

    _set_cell_paragraph_lines(left1, ["Số: 1782/CV-ĐV", "V/v phúc đáp Công văn số 1782/TTKSBT-PCBTN"])
    for p in left1.paragraphs:
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        for run in p.runs:
            _set_run_font_tnr(run, size_pt=13, bold=False, italic=False)

    _set_cell_paragraph_lines(right1, ["Quảng Ninh, ngày 02 tháng 10 năm 2025"])
    right1.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.RIGHT
    for run in right1.paragraphs[0].runs:
        _set_run_font_tnr(run, size_pt=13, bold=False, italic=True)

    # Spacer row
    c20 = t0.rows[2].cells[0]
    c21 = t0.rows[2].cells[1]
    c20.merge(c21)
    c20.paragraphs[0].text = " "

    ensure_underline_after_agency_cell(left0)
    ensure_underline_after_motto_cell(right0)

    # Body paragraphs
    for t in texts:
        p = doc.add_paragraph(t)
        if justify_body and t.strip():
            p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        p.paragraph_format.first_line_indent = Cm(1.27)
        for run in p.runs:
            _set_run_font_tnr(run, size_pt=body_pt)

    # Exactly one blank line between body content and signature/distribution table.
    gap = doc.add_paragraph("")
    gap.paragraph_format.space_before = Pt(0)
    gap.paragraph_format.space_after = Pt(0)

    # Footer block is in document body (not Word footer).
    t1 = doc.add_table(rows=1, cols=2)
    _table_no_borders(t1)
    _set_table_fixed_layout(t1, TOTAL_CONTENT_WIDTH_TWIPS)
    t1.autofit = False
    _set_cell_width(t1.rows[0].cells[0], LEFT_COL_TWIPS)
    _set_cell_width(t1.rows[0].cells[1], RIGHT_COL_TWIPS)

    noi_nhan = t1.rows[0].cells[0]
    noi_nhan.paragraphs[0].text = "Nơi nhận:"
    noi_nhan.add_paragraph("- Như trên;")
    noi_nhan.add_paragraph("- Lưu: VT, VP.")
    for i, p in enumerate(noi_nhan.paragraphs):
        p.alignment = WD_ALIGN_PARAGRAPH.LEFT
        for run in p.runs:
            if i == 0:
                _set_run_font_tnr(run, size_pt=12, bold=True, italic=True)
            else:
                _set_run_font_tnr(run, size_pt=11, bold=False, italic=False)

    signer = t1.rows[0].cells[1]
    _set_para_with_br(signer.paragraphs[0], "KT. GIÁM ĐỐC\nPHÓ GIÁM ĐỐC")
    signer.add_paragraph("")
    signer.add_paragraph("")
    signer.add_paragraph("")
    signer.add_paragraph("Nguyễn Văn A")
    for p in signer.paragraphs:
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        for run in p.runs:
            _set_run_font_tnr(run, size_pt=13, bold=True, italic=False)

    _set_spacing_zero_everywhere(doc)
    return doc


def _placeholder_value(token: str) -> str:
    key = token.strip().lower()
    if "tên cơ quan chủ quản" in key or "ten co quan chu quan" in key:
        return "SỞ Y TẾ TỈNH QUẢNG NINH"
    if "tên cơ quan ban hành" in key or "ten co quan ban hanh" in key:
        return "TRUNG TÂM Y TẾ ĐỊA PHƯƠNG"
    if "tên đơn vị" in key or "ten don vi" in key:
        return "Trung tâm Y tế địa phương"
    if "địa danh" in key or "dia danh" in key:
        return "Quảng Ninh"
    if "đơn vị cấp trên" in key or "don vi cap tren" in key:
        return "Sở Y tế tỉnh Quảng Ninh"
    if "bộ phận" in key or "bo phan" in key:
        return "CNTT"
    if "họ và tên" in key or "ho va ten" in key:
        return "Nguyễn Văn A"
    return token.strip()


def _strip_square_placeholders(text: str) -> str:
    return re.sub(r"\[([^\]]+)\]", lambda m: _placeholder_value(m.group(1)), text)


def _extract_source_body_paragraph_texts(source: Document) -> list[str]:
    texts: list[str] = []
    for p in source.paragraphs:
        t = _strip_square_placeholders(p.text.strip())
        if not t:
            continue
        texts.append(t)
    return texts


def _has_vietnamese_diacritics(texts: list[str]) -> bool:
    merged = "\n".join(texts).lower()
    return any(ch in VIETNAMESE_DIACRITICS for ch in merged)


def _count_suspect_lossy_chars(text: str) -> int:
    # Documents with many '?' or replacement chars are usually mojibake/lossy text.
    return text.count("?") + text.count("�")


def _assert_doc_not_lossy(doc: Document, *, context: str) -> None:
    texts: list[str] = [p.text for p in doc.paragraphs]
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                texts.extend(p.text for p in cell.paragraphs)
    merged = "\n".join(texts)
    if _count_suspect_lossy_chars(merged) >= 5:
        raise SystemExit(
            f"{context}: Phát hiện ký tự '?' hoặc '�' bất thường trong văn bản. "
            "Dừng xử lý để tránh làm mất dấu tiếng Việt. Hãy dùng nguồn có dấu hoặc sửa nội dung trước."
        )


def _run_has_drawing(run) -> bool:
    """Return True if the run element contains a <w:drawing> — must not touch .text."""
    return run._r.find(qn("w:drawing")) is not None or run._r.find(qn("w:pict")) is not None


def _sanitize_all_placeholders(doc: Document) -> None:
    for p in doc.paragraphs:
        for run in p.runs:
            if not _run_has_drawing(run):
                run.text = _strip_square_placeholders(run.text)
    for tbl in doc.tables:
        for row in tbl.rows:
            for cell in row.cells:
                for p in cell.paragraphs:
                    for run in p.runs:
                        if not _run_has_drawing(run):
                            run.text = _strip_square_placeholders(run.text)


def _fill_default_header_fields(doc: Document) -> None:
    if len(doc.tables) < 2:
        return
    t0 = doc.tables[0]
    left0 = t0.rows[0].cells[0]
    left1 = t0.rows[1].cells[0]
    right1 = t0.rows[1].cells[-1]

    def _replace_cell_text(cell, replacements: dict[str, str]) -> None:
        for p in cell.paragraphs:
            for run in p.runs:
                if run._r.find(qn("w:drawing")) is not None or run._r.find(qn("w:pict")) is not None:
                    continue  # never overwrite runs that hold inline shapes
                txt = run.text
                for old, new in replacements.items():
                    txt = txt.replace(old, new)
                run.text = txt

    _replace_cell_text(
        left0,
        {
            "[TÊN CƠ QUAN CHỦ QUẢN]": "SỞ Y TẾ TỈNH QUẢNG NINH",
            "[TÊN ĐƠN VỊ]": "TRUNG TÂM Y TẾ ĐỊA PHƯƠNG",
        },
    )
    _replace_cell_text(
        left1,
        {
            "Số: …/…": "Số: 1782/CV-ĐV",
            "V/v … (trích yếu công việc)": "V/v phúc đáp Công văn số 1782/TTKSBT-PCBTN",
            "Số: .../...-...": "Số: 1782/CV-ĐV",
        },
    )
    _replace_cell_text(
        right1,
        {
            "[Địa danh], ngày ... tháng ... năm 2025": "Quảng Ninh, ngày 02 tháng 10 năm 2025",
            "……, ngày … tháng … năm 2026": "Quảng Ninh, ngày 02 tháng 10 năm 2025",
        },
    )


def _insert_paragraphs_before_element(doc: Document, before_el, texts: list[str], *, justify: bool, body_pt: int | None) -> None:
    body = doc._body._body  # type: ignore[attr-defined]
    insert_at = list(body.iterchildren()).index(before_el)
    new_paragraph_els = []
    for t in texts:
        p = doc.add_paragraph(t)
        if justify and t.strip():
            p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        p.paragraph_format.space_before = Pt(0)
        p.paragraph_format.space_after = Pt(0)
        p.paragraph_format.first_line_indent = Cm(1.27)
        for run in p.runs:
            _set_run_font_tnr(run, size_pt=body_pt)
        new_paragraph_els.append(p._element)
    for p_el in new_paragraph_els:
        body.remove(p_el)
    for p_el in new_paragraph_els:
        body.insert(insert_at, p_el)
        insert_at += 1


def _ensure_single_blank_line_before_signature_table(doc: Document) -> None:
    if len(doc.tables) < 2:
        return
    body = doc._body._body  # type: ignore[attr-defined]
    sign_tbl_el = doc.tables[1]._tbl
    children = list(body.iterchildren())
    try:
        i_sign = children.index(sign_tbl_el)
    except ValueError:
        return
    if i_sign <= 0:
        return
    i = i_sign - 1
    while i >= 0:
        el = children[i]
        if el.tag != qn("w:p"):
            break
        txt = "".join(el.itertext()).strip()
        if txt:
            break
        body.remove(el)
        children.pop(i)
        i -= 1
    blank = doc.add_paragraph("")
    blank.paragraph_format.space_before = Pt(0)
    blank.paragraph_format.space_after = Pt(0)
    blank_el = blank._element
    body.remove(blank_el)
    children = list(body.iterchildren())
    i_sign = children.index(sign_tbl_el)
    body.insert(i_sign, blank_el)


def _format_noi_nhan_block(doc: Document) -> None:
    if len(doc.tables) < 2:
        return
    left = doc.tables[1].rows[0].cells[0]

    # Normalize cases where all "Nơi nhận" lines are packed in one paragraph.
    if len(left.paragraphs) == 1 and "\n" in left.paragraphs[0].text:
        lines = [ln.strip() for ln in left.paragraphs[0].text.splitlines() if ln.strip()]
        left.paragraphs[0].text = lines[0] if lines else "Nơi nhận:"
        for ln in lines[1:]:
            left.add_paragraph(ln)

    non_empty_idx = 0
    for p in left.paragraphs:
        if not p.text.strip():
            continue
        p.alignment = WD_ALIGN_PARAGRAPH.LEFT
        p.paragraph_format.space_before = Pt(0)
        p.paragraph_format.space_after = Pt(0)
        for run in p.runs:
            if non_empty_idx == 0:
                _set_run_font_tnr(run, size_pt=12, bold=True, italic=True)
                run.font.underline = False
            else:
                _set_run_font_tnr(run, size_pt=11, bold=False, italic=False)
                run.font.underline = False
        non_empty_idx += 1


def _set_spacing_zero_everywhere(doc: Document) -> None:
    def _force_spacing_xml_zero(paragraph) -> None:
        ppr = paragraph._p.get_or_add_pPr()
        # Find or create w:spacing, preserving existing w:line / w:lineRule.
        sp = ppr.find(qn("w:spacing"))
        if sp is None:
            sp = OxmlElement("w:spacing")
            ppr.append(sp)
        sp.set(qn("w:before"), "0")
        sp.set(qn("w:after"), "0")
        # Remove autospacing flags that override before/after=0.
        sp.attrib.pop(qn("w:beforeAutospacing"), None)
        sp.attrib.pop(qn("w:afterAutospacing"), None)

    # 1) Force all paragraph styles to 0/0 to avoid Word default "After 10 pt".
    try:
        for style in doc.styles:
            if style.type != WD_STYLE_TYPE.PARAGRAPH:
                continue
            pf = style.paragraph_format
            pf.space_before = Pt(0)
            pf.space_after = Pt(0)
    except Exception:
        # Some documents may contain protected/unexpected styles.
        pass

    # 2) Force concrete paragraphs in document body.
    for p in doc.paragraphs:
        p.paragraph_format.space_before = Pt(0)
        p.paragraph_format.space_after = Pt(0)
        _force_spacing_xml_zero(p)

    # 3) Force concrete paragraphs inside all table cells.
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for p in cell.paragraphs:
                    p.paragraph_format.space_before = Pt(0)
                    p.paragraph_format.space_after = Pt(0)
                    _force_spacing_xml_zero(p)

    # 4) Force header/footer paragraphs for all sections.
    for section in doc.sections:
        for p in section.header.paragraphs:
            p.paragraph_format.space_before = Pt(0)
            p.paragraph_format.space_after = Pt(0)
            _force_spacing_xml_zero(p)
        for table in section.header.tables:
            for row in table.rows:
                for cell in row.cells:
                    for p in cell.paragraphs:
                        p.paragraph_format.space_before = Pt(0)
                        p.paragraph_format.space_after = Pt(0)
                        _force_spacing_xml_zero(p)
        for p in section.footer.paragraphs:
            p.paragraph_format.space_before = Pt(0)
            p.paragraph_format.space_after = Pt(0)
            _force_spacing_xml_zero(p)
        for table in section.footer.tables:
            for row in table.rows:
                for cell in row.cells:
                    for p in cell.paragraphs:
                        p.paragraph_format.space_before = Pt(0)
                        p.paragraph_format.space_after = Pt(0)
                        _force_spacing_xml_zero(p)


def _ensure_agency_header_lines(left0_cell) -> None:
    lines = [p.text.strip() for p in left0_cell.paragraphs if p.text.strip()]
    if not lines:
        lines = ["SỞ Y TẾ TỈNH QUẢNG NINH", "BỆNH VIỆN ĐA KHOA TỈNH QUẢNG NINH"]
    elif len(lines) == 1:
        lines = ["SỞ Y TẾ TỈNH QUẢNG NINH", lines[0]]

    # Rule:
    # - cơ quan chủ quản: đúng 1 dòng (không chứa xuống dòng)
    # - tên đơn vị: tối đa 2 dòng
    manager_line = " ".join(lines[0].replace("\n", " ").split())
    unit_text = " ".join(" ".join(lines[1:]).replace("\n", " ").split())
    if not unit_text:
        unit_text = "BỆNH VIỆN ĐA KHOA TỈNH QUẢNG NINH"
    unit_parts = [x.strip() for x in unit_text.split(" / ") if x.strip()]
    if len(unit_parts) >= 2:
        unit_line = f"{unit_parts[0]}\n{' / '.join(unit_parts[1:])}"
    else:
        # Soft split for long names: keep at most one line break.
        words = unit_text.split()
        if len(words) > 8:
            cut = len(words) // 2
            unit_line = f"{' '.join(words[:cut])}\n{' '.join(words[cut:])}"
        else:
            unit_line = unit_text
    # Cap to maximum 2 lines for unit block.
    unit_lines = [ln.strip() for ln in unit_line.splitlines() if ln.strip()]
    if len(unit_lines) > 2:
        unit_line = f"{unit_lines[0]}\n{' '.join(unit_lines[1:])}"
    else:
        unit_line = "\n".join(unit_lines)

    while len(left0_cell.paragraphs) < 2:
        left0_cell.add_paragraph("")

    left0_cell.paragraphs[0].text = manager_line
    _set_para_with_br(left0_cell.paragraphs[1], unit_line)

    # Keep only 2 semantic lines before the underline shape.
    while len(left0_cell.paragraphs) > 2:
        el = left0_cell.paragraphs[-1]._element
        el.getparent().remove(el)

    p0, p1 = left0_cell.paragraphs[0], left0_cell.paragraphs[1]
    p0.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p1.alignment = WD_ALIGN_PARAGRAPH.CENTER
    # Keep manager line in one visual line when text is long.
    manager_size = 11 if len(manager_line) > 34 else 12 if len(manager_line) > 28 else 13
    for run in p0.runs:
        _set_run_font_tnr(run, size_pt=manager_size, bold=False, italic=False)
    for run in p1.runs:
        _set_run_font_tnr(run, size_pt=13, bold=True, italic=False)


def _fix_table0(doc: Document) -> None:
    t0 = doc.tables[0]

    def _lr(row_idx: int):
        row = t0.rows[row_idx]
        return row.cells[0], row.cells[-1]

    for ri in (0, 1):
        left, right = _lr(ri)
        _set_cell_width(left, LEFT_COL_TWIPS)
        _set_cell_width(right, RIGHT_COL_TWIPS)

    left0, right0 = _lr(0)
    for p in left0.paragraphs:
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    # Ensure agency unit line (2nd paragraph) is bold.
    for idx, p in enumerate(left0.paragraphs):
        for run in p.runs:
            if idx == 0:
                _set_run_font_tnr(run, size_pt=13, bold=False, italic=False)
            else:
                _set_run_font_tnr(run, size_pt=13, bold=True, italic=False)
    for p in right0.paragraphs:
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER

    left1, right1 = _lr(1)
    for p in left1.paragraphs:
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    for p in right1.paragraphs:
        p.alignment = WD_ALIGN_PARAGRAPH.RIGHT

    # Preserve original text content; only normalize alignment/font and lines.
    if right0.paragraphs:
        p_qh = right0.paragraphs[0]
        p_qh.alignment = WD_ALIGN_PARAGRAPH.CENTER
        for run in p_qh.runs:
            _set_run_font_tnr(run, size_pt=13, bold=True, italic=False)
    ensure_underline_after_agency_cell(left0)
    ensure_underline_after_motto_cell(right0)


def _remove_extra_paragraphs(cell, keep_first_n: int) -> None:
    while len(cell.paragraphs) > keep_first_n:
        el = cell.paragraphs[-1]._element
        el.getparent().remove(el)


def _fix_table1_signature_cell(doc: Document) -> None:
    right = doc.tables[1].rows[0].cells[1]
    # Preserve original text; only enforce formatting.
    last_non_empty = None
    for p in right.paragraphs:
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        if p.text.strip():
            last_non_empty = p
        for run in p.runs:
            _set_run_font_tnr(run, size_pt=13, bold=True, italic=False)
    if last_non_empty is not None:
        for run in last_non_empty.runs:
            _set_run_font_tnr(run, size_pt=13, bold=True, italic=False)
    _format_noi_nhan_block(doc)


def _justify_main_body_paragraphs(doc: Document) -> None:
    for p in doc.paragraphs:
        if p.text.strip():
            p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
            p.paragraph_format.space_before = Pt(0)
            p.paragraph_format.space_after = Pt(0)
            p.paragraph_format.first_line_indent = Cm(1.27)


def _set_main_body_font_size(doc: Document, size_pt: int) -> None:
    for p in doc.paragraphs:
        if not p.text.strip():
            continue
        for run in p.runs:
            _set_run_font_tnr(run, size_pt=size_pt, bold=run.bold, italic=run.italic)


def _remove_blank_body_paragraphs(doc: Document) -> None:
    for p in list(doc.paragraphs):
        if p.text.strip():
            continue
        el = p._element
        el.getparent().remove(el)


def _add_page_field(run) -> None:
    fld_char_begin = OxmlElement("w:fldChar")
    fld_char_begin.set(qn("w:fldCharType"), "begin")
    instr = OxmlElement("w:instrText")
    instr.set(qn("xml:space"), "preserve")
    instr.text = "PAGE"
    fld_char_end = OxmlElement("w:fldChar")
    fld_char_end.set(qn("w:fldCharType"), "end")
    run._r.append(fld_char_begin)
    run._r.append(instr)
    run._r.append(fld_char_end)


def _configure_normal_style(document: Document, body_pt: int) -> None:
    normal = document.styles["Normal"]
    normal.font.name = "Times New Roman"
    normal.font.size = Pt(body_pt)
    rpr = normal.element.get_or_add_rPr()
    rfonts = rpr.get_or_add_rFonts()
    rfonts.set(qn("w:ascii"), "Times New Roman")
    rfonts.set(qn("w:hAnsi"), "Times New Roman")
    rfonts.set(qn("w:eastAsia"), "Times New Roman")


def _set_section_a4_margins(section) -> None:
    section.page_height = Mm(297)
    section.page_width = Mm(210)
    section.top_margin = Mm(20)
    section.bottom_margin = Mm(20)
    section.left_margin = Mm(30)
    section.right_margin = Mm(15)


def _paragraph_body_format(p, *, line_pt: int | None) -> None:
    pf = p.paragraph_format
    pf.space_before = Pt(0)
    pf.space_after = Pt(0)
    pf.first_line_indent = Cm(1.27)
    if line_pt is None:
        pf.line_spacing_rule = WD_LINE_SPACING.SINGLE
    else:
        pf.line_spacing_rule = WD_LINE_SPACING.EXACTLY
        pf.line_spacing = Pt(line_pt)


def cmd_rebuild(args) -> None:
    source_path = Path(args.source)
    output_path = Path(args.output)
    src_doc = Document(str(source_path))
    _assert_doc_not_lossy(src_doc, context="Source")
    texts = _extract_source_body_paragraph_texts(src_doc)
    if texts and not _has_vietnamese_diacritics(texts):
        raise SystemExit(
            "Nguon noi dung co dau hieu mat dau tieng Viet. "
            "Dung xuat file de tranh vo dinh dang/noi dung. "
            "Hay cung cap ban noi dung co dau hoac file goc co dau."
        )
    out_doc = _build_nd30_document_from_source(texts, justify_body=not args.no_justify_body, body_pt=args.body_pt)
    _sanitize_all_placeholders(out_doc)
    _assert_header_underlines(out_doc)
    if output_path.exists():
        output_path.unlink()
    out_doc.save(str(output_path))
    print("Saved")


def cmd_fix(args) -> None:
    path = Path(args.docx)
    doc = Document(str(path))
    suspect_before = _doc_suspect_count(doc)
    if args.spacing_only:
        _set_spacing_zero_everywhere(doc)
        if _doc_suspect_count(doc) > suspect_before:
            raise SystemExit("Dừng lưu: phát sinh thêm ký tự '?'/'�' sau khi xử lý spacing.")
        doc.save(str(path))
        print("Saved spacing-only (before/after = 0pt).")
        return

    _assert_doc_not_lossy(doc, context="Input")
    if not args.apply_layout:
        print("Vietnamese check passed. No layout changes applied.")
        return
    _fix_table0(doc)
    _fix_table1_signature_cell(doc)
    if not args.no_justify_body:
        _justify_main_body_paragraphs(doc)
    _set_main_body_font_size(doc, args.body_pt)
    if not args.keep_empty_lines:
        _remove_blank_body_paragraphs(doc)
    _ensure_single_blank_line_before_signature_table(doc)
    _set_spacing_zero_everywhere(doc)
    ensure_underline_after_agency_cell(doc.tables[0].rows[0].cells[0])
    ensure_underline_after_motto_cell(doc.tables[0].rows[0].cells[-1])
    _assert_header_underlines(doc)
    if _doc_suspect_count(doc) > suspect_before:
        raise SystemExit("Dừng lưu: phát sinh thêm ký tự '?'/'�' sau khi apply layout.")
    doc.save(str(path))
    print("Saved")


def cmd_legacy(args) -> None:
    if not args.allow_legacy_stack:
        raise SystemExit(
            "Refuse to generate legacy centered-stack output. "
            "Use Mau_cong_van_ND30_tai_ve.docx + scripts/office_skill_cli.py rebuild for ND30-compliant files. "
            "If you really need legacy stack, rerun with --allow-legacy-stack."
        )
    out = Path(args.output)
    if out.suffix.lower() != ".docx":
        raise SystemExit("--output phai la file .docx")
    noi_dung = ""
    if args.noi_dung_file:
        noi_dung = Path(args.noi_dung_file).read_text(encoding="utf-8")
    line_exact = None if args.line_exact_pt == 0 else args.line_exact_pt

    document = Document()
    _configure_normal_style(document, args.body_pt)
    section = document.sections[0]
    _set_section_a4_margins(section)
    header = section.header
    hp = header.paragraphs[0] if header.paragraphs else header.add_paragraph()
    hp.text = ""
    hp.alignment = WD_ALIGN_PARAGRAPH.CENTER
    hr = hp.add_run()
    _set_run_font_tnr(hr, size_pt=args.body_pt, bold=False, italic=False)
    _add_page_field(hr)

    p = document.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    _set_run_font_tnr(p.add_run("CỘNG HÒA XÃ HỘI CHỦ NGHĨA VIỆT NAM"), size_pt=13, bold=True, italic=False)
    p = document.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    _set_run_font_tnr(p.add_run("Độc lập - Tự do - Hạnh phúc"), size_pt=13, bold=True, italic=False)
    document.add_paragraph()
    p = document.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    _set_run_font_tnr(p.add_run(args.co_quan.strip().upper()), size_pt=13, bold=True, italic=False)
    p = document.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    _set_run_font_tnr(p.add_run(args.so_ky_hieu.strip()), size_pt=13, bold=False, italic=False)
    p = document.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    _set_run_font_tnr(p.add_run(args.dia_danh_ngay.strip()), size_pt=13, bold=False, italic=True)
    document.add_paragraph()
    p = document.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    _set_run_font_tnr(p.add_run(args.trich_yeu.strip()), size_pt=13, bold=False, italic=False)
    _paragraph_body_format(p, line_pt=line_exact)
    document.add_paragraph()
    for line in noi_dung.splitlines():
        p = document.add_paragraph(line)
        p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        _paragraph_body_format(p, line_pt=line_exact)
        for run in p.runs:
            _set_run_font_tnr(run, size_pt=args.body_pt, bold=False, italic=False)
    _set_spacing_zero_everywhere(document)
    document.save(str(out))
    print("Saved")


def main() -> None:
    parser = argparse.ArgumentParser(description="Office skill ND30 utilities.")
    sub = parser.add_subparsers(dest="cmd", required=True)

    p_rebuild = sub.add_parser("rebuild", help="Rebuild from frozen ND30 template.")
    p_rebuild.add_argument("--source", required=True)
    p_rebuild.add_argument("--output", required=True)
    p_rebuild.add_argument("--no-justify-body", action="store_true")
    p_rebuild.add_argument("--body-pt", type=int, choices=(13, 14), default=14)
    p_rebuild.set_defaults(func=cmd_rebuild)

    p_fix = sub.add_parser("fix", help="Fix layout for existing ND30 doc.")
    p_fix.add_argument("docx")
    p_fix.add_argument(
        "--spacing-only",
        action="store_true",
        help="Only force paragraph spacing before/after to 0pt; do not touch text, lines, or layout.",
    )
    p_fix.add_argument(
        "--apply-layout",
        action="store_true",
        help="Apply layout/format rules. Default only checks Vietnamese integrity and keeps formatting unchanged.",
    )
    p_fix.add_argument("--no-justify-body", action="store_true")
    p_fix.add_argument("--body-pt", type=int, choices=(13, 14), default=14)
    p_fix.add_argument("--keep-empty-lines", action="store_true")
    p_fix.set_defaults(func=cmd_fix)

    p_legacy = sub.add_parser("legacy", help="Legacy centered-stack generator.")
    p_legacy.add_argument("--allow-legacy-stack", action="store_true")
    p_legacy.add_argument("--output", required=True)
    p_legacy.add_argument("--co-quan", default="TÊN CƠ QUAN BAN HÀNH")
    p_legacy.add_argument("--so-ky-hieu", default=".../...")
    p_legacy.add_argument("--dia-danh-ngay", default="……, ngày … tháng … năm …")
    p_legacy.add_argument("--trich-yeu", default="V/v …")
    p_legacy.add_argument("--noi-dung-file", default="")
    p_legacy.add_argument("--body-pt", type=int, choices=(13, 14), default=14)
    p_legacy.add_argument("--line-exact-pt", type=int, default=16)
    p_legacy.set_defaults(func=cmd_legacy)

    args = parser.parse_args()
    try:
        sys.stdout.reconfigure(encoding="utf-8")  # type: ignore[attr-defined]
    except Exception:
        pass
    try:
        sys.stderr.reconfigure(encoding="utf-8")  # type: ignore[attr-defined]
    except Exception:
        pass
    args.func(args)


if __name__ == "__main__":
    main()